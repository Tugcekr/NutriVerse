"""
Created on Sat Oct 25 04:08:03 2025

@author: Tugce
"""

import docx
from langchain.llms import Ollama
from langchain.chains import LLMChain
from langchain.prompts import PromptTemplate
from langchain.embeddings import HuggingFaceEmbeddings
from langchain.vectorstores import FAISS
from langchain.text_splitter import CharacterTextSplitter
from langchain.schema import Document

class RAGAnalysisAgent:
    def __init__(self, docx_path):
        self.knowledge_base = ""
        self.vector_store = None
        self.llm = Ollama(model="llama3.2:3b")
        self.load_hazard_data_from_docx(docx_path)
        self.setup_vector_database() 
        self.setup_prompts()
    
    def setup_vector_database(self):
        text_splitter = CharacterTextSplitter(
            chunk_size=500,chunk_overlap=50,separator="\n")
        
        texts = text_splitter.split_text(self.knowledge_base)
        documents = [Document(page_content=text) for text in texts]
        embeddings = HuggingFaceEmbeddings(
            model_name="sentence-transformers/all-MiniLM-L6-v2")
        self.vector_store = FAISS.from_documents(documents, embeddings)
        print("Vector database created!")
    
    def setup_prompts(self):
        self.analysis_prompt = PromptTemplate(
            input_variables=["ingredients", "diet", "hazard_info"],
            template="""
            You are a food safety and nutrition expert. Analyze the product ingredients for specific dietary requirements.

            PRODUCT INGREDIENTS: {ingredients}
            DIETARY REQUIREMENT: {diet}
            HAZARD KNOWLEDGE: {hazard_info}

            Provide analysis in this exact format:
            SUITABLE: [YES/NO]
            RISK_LEVEL: [LOW/MEDIUM/HIGH]
            HAZARDOUS_INGREDIENTS: [comma-separated list or "None"]
            EXPLANATION: [Clear explanation based on scientific evidence]

            Focus on:
            - Scientific evidence from hazard knowledge
            - Health impacts and risks
            - Specific ingredients that violate dietary requirements
            - Safety recommendations

            Be precise and evidence-based in your analysis.
            """)
    
    def load_hazard_data_from_docx(self, docx_path):
        doc = docx.Document(docx_path)
        full_text = []
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                full_text.append(paragraph.text)
        
        self.knowledge_base = "\n".join(full_text)
        print(f" Document loaded: {len(self.knowledge_base)} characters")
    
    def extract_diet_info(self, diet_keyword):
        if self.vector_store is None:
            return "Vector database not ready"
        diet_queries = {
            'celiac': "gluten wheat barley rye celiac disease autoimmune",
            'diabetes': "sugar glucose fructose carbohydrates glycemic insulin",
            'vegan': "vegan animal milk egg honey gelatin dairy",
            'vegetarian': "vegetarian meat fish chicken poultry gelatin",
            'lactose': "lactose milk dairy cheese whey intolerance",
            'nut_allergy': "nut peanut almond walnut hazelnut allergy anaphylaxis",
            'soy_allergy': "soy soybean tofu soybeans allergy",
            'heart_disease': "saturated fat cholesterol sodium salt heart cardiovascular",
            'hypertension': "sodium salt blood pressure hypertension",
            'baby_0_6': "infant formula breast milk 0-6 months honey salt sugar",
            'baby_6_8': "6-8 months puree salt sugar honey egg whites",
            'baby_8_12': "8-12 months soft foods choking hazards salt sugar"
        }
        
        if diet_keyword in diet_queries:
            docs = self.vector_store.similarity_search(
                diet_queries[diet_keyword], k=3)
            relevant_info = "\n".join([doc.page_content for doc in docs])
            return relevant_info if relevant_info else "No relevant information found"
        return "Diet information not found"
    
    def analyze_with_llm(self, ingredients, diet):
        hazard_info = self.extract_diet_info(diet)
        chain = LLMChain(llm=self.llm, prompt=self.analysis_prompt)
        try:
            response = chain.run({
                "ingredients": ingredients,
                "diet": diet,
                "hazard_info": hazard_info})
            return self.parse_llm_response(response, diet)
        except Exception as e:
            return {
                'suitable': False,
                'risk_level': 'HIGH',
                'explanation': f'LLM error: {str(e)}',
                'hazardous_ingredients': ['Analysis error']
            }
    
    def parse_llm_response(self, response, diet):
        response_lower = response.lower()
        suitable = False
        risk_level = 'HIGH'
        explanation = "Analysis incomplete"
        hazardous_ingredients = []
        if 'suitable: yes' in response_lower or 'suitable:yes' in response_lower:
            suitable = True
        elif 'suitable: no' in response_lower or 'suitable:no' in response_lower:
            suitable = False
        if 'risk_level: low' in response_lower:
            risk_level = 'LOW'
        elif 'risk_level: medium' in response_lower:
            risk_level = 'MEDIUM'
        elif 'risk_level: high' in response_lower:
            risk_level = 'HIGH'
        if 'hazardous_ingredients:' in response_lower:
            start_idx = response_lower.find('hazardous_ingredients:') + len('hazardous_ingredients:')
            end_idx = response_lower.find('explanation:', start_idx)
            if end_idx == -1:
                end_idx = len(response)
            
            ingredients_text = response[start_idx:end_idx].strip()
            if 'none' not in ingredients_text.lower():
                hazardous_ingredients = [ing.strip() for ing in ingredients_text.split(',') if ing.strip()]
        
        if 'explanation:' in response_lower:
            start_idx = response_lower.find('explanation:') + len('explanation:')
            explanation = response[start_idx:].strip()
        
        diet_hazards = {
            'celiac': ['wheat', 'gluten', 'barley', 'rye', 'oats', 'malt'],
            'diabetes': ['sugar', 'glucose', 'fructose', 'sucrose', 'syrup', 'honey'],
            'vegan': ['milk', 'egg', 'honey', 'gelatin', 'cheese', 'yogurt'],
            'vegetarian': ['meat', 'fish', 'chicken', 'beef', 'pork', 'gelatin'],
            'lactose': ['milk', 'cheese', 'yogurt', 'whey', 'lactose', 'dairy'],
            'nut_allergy': ['peanut', 'almond', 'walnut', 'hazelnut', 'cashew', 'pistachio'],
            'soy_allergy': ['soy', 'soybean', 'tofu', 'soy lecithin', 'soy protein'],
            'heart_disease': ['saturated fat', 'trans fat', 'cholesterol', 'sodium', 'salt'],
            'hypertension': ['sodium', 'salt', 'msg', 'monosodium glutamate'],
            'baby_0_6': ['honey', 'salt', 'sugar', 'cow milk'],
            'baby_6_8': ['honey', 'salt', 'sugar', 'egg whites', 'nuts'],
            'baby_8_12': ['honey', 'salt', 'sugar', 'choking hazards']
        }
        
        if not hazardous_ingredients and diet in diet_hazards:
            ingredients_lower = response_lower
            for hazard in diet_hazards[diet]:
                if hazard in ingredients_lower:
                    hazardous_ingredients.append(hazard)
        
        return {
            'suitable': suitable,
            'risk_level': risk_level,
            'explanation': explanation,
            'hazardous_ingredients': hazardous_ingredients if hazardous_ingredients else ['No hazardous ingredients detected']
        }
    
    def analyze_ingredients(self, ingredients_text, user_preferences):
        analysis_results = {}
        
        for preference in user_preferences:
            result = self.analyze_with_llm(ingredients_text, preference)
            analysis_results[preference] = result
        
        return analysis_results
    
    def calculate_risk_score(self, analysis_results):
        total_preferences = len(analysis_results)
        if total_preferences == 0:
            return 0
        
        safe_preferences = sum(1 for result in analysis_results.values() if result['suitable'])
        return (safe_preferences / total_preferences) * 100

    def search_similar_products(self, query, k=3):
        if self.vector_store:
            docs = self.vector_store.similarity_search(query, k=k)
            return [doc.page_content for doc in docs]
        return []

    def get_scientific_evidence(self, ingredient, diet):
        try:
            query = f"{ingredient} {diet} health effects scientific research"
            docs = self.vector_store.similarity_search(query, k=2)
            evidence = "\n".join([doc.page_content for doc in docs])
            return evidence if evidence else "No specific scientific evidence found"
        except:
            return "Evidence search unavailable"

    def analyze_for_baby_age(self, ingredients, baby_age_months):
        age_groups = {
            (0, 6): 'baby_0_6',
            (6, 8): 'baby_6_8', 
            (8, 12): 'baby_8_12' }
        
        for (min_age, max_age), diet_key in age_groups.items():
            if min_age <= baby_age_months <= max_age:
                return self.analyze_with_llm(ingredients, diet_key)
        
        return {
            'suitable': True,
            'risk_level': 'LOW',
            'explanation': 'General food safety guidelines apply',
            'hazardous_ingredients': ['None']
        }

# Usage example
if __name__ == "__main__":
    agent = RAGAnalysisAgent("C:/Users/Tugce/OneDrive/Masaüstü/hazard_ingredients_short.docx")
    
    # Test analysis
    test_ingredients = "wheat flour, milk powder, egg, soy lecithin, sugar, salt"
    user_diets = ['celiac', 'diabetes', 'lactose']
    
    print("Starting analysis...")
    results = agent.analyze_ingredients(test_ingredients, user_diets)
    risk_score = agent.calculate_risk_score(results)
    
    print(f"\nProduct Ingredients: {test_ingredients}")
    print(f"Overall Risk Score: {risk_score:.1f}%")
    
    for diet, analysis in results.items():
        print(f"\n{diet.upper()}:")
        print(f"  Suitable: {'Yes' if analysis['suitable'] else 'No'}")
        print(f"  Risk Level: {analysis['risk_level']}")
        print(f"  Explanation: {analysis['explanation']}")
        if analysis['hazardous_ingredients']:
            print(f"  Hazardous Ingredients: {', '.join(analysis['hazardous_ingredients'])}")
    
    print("\n Semantic Search Test:")
    similar = agent.search_similar_products("gluten free products", 2)
    for i, doc in enumerate(similar):
        print(f"  {i+1}. {doc[:100]}...")
    
    print("\n Scientific Evidence Test:")
    evidence = agent.get_scientific_evidence("gluten", "celiac")
    print(f"Evidence: {evidence[:200]}...")
    
    print("\n Baby Safety Test:")
    baby_analysis = agent.analyze_for_baby_age("honey, salt, cow milk", 4)
    print(f"Baby Safety: {baby_analysis}")